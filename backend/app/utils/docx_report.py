"""
DOCX report generation — app.utils.docx_report

Renders the same normalized export dict used by app.utils.pdf_report
(see app.utils.report_export) into a fully formatted Word document
using python-docx. Section order and content mirror the PDF exactly so
the two formats never drift apart.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD_DARK = RGBColor(0x6B, 0x5A, 0x34)
SLATE = RGBColor(0x1C, 0x24, 0x33)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)

PIPELINE_STAGES = [
    "Case Intake", "Legal Research", "Document Intelligence",
    "Compliance Review", "Red Flag Detection", "Risk Assessment", "Legal Report",
]

DISCLAIMER_TEXT = (
    "This report is AI-generated and is not a substitute for professional legal advice. "
    "Consult a licensed attorney before making decisions based on this analysis. "
    "Findings are observations grounded in the evidence cited and do not constitute "
    "legal conclusions or recommendations to act."
)


def _shade_cell(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _heading(doc, text: str):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = GOLD_DARK
        run.font.size = Pt(14)


def _table(doc, header: list, rows: list):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        _shade_cell(hdr[i], "1C2433")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return t


def generate_docx_report(data: dict) -> bytes:
    doc = Document()

    # ---- Branding header ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run("LEGAL INTELLIGENCE COMMAND ROOM")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GOLD_DARK

    case_title = doc.add_heading("Legal Intelligence Report", level=0)
    for r in case_title.runs:
        r.font.color.rgb = SLATE

    # ---- 1. Case Information ----
    _heading(doc, "1. Case Information")
    case_rows = [
        ["Case ID", data["case_id"]],
        ["Task Type", data.get("task_type", "—")],
        ["Workflow Mode", (data.get("workflow_mode") or "single").upper()],
        ["Query", data.get("user_query", "—")],
        ["Jurisdiction", data.get("jurisdiction") or "—"],
        ["Governing Law", data.get("governing_law") or "—"],
        ["Generated On", data.get("generated_on", "—")],
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light List Accent 1"
    for label, value in case_rows:
        row = t.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.bold = True
        row[1].text = str(value)
    doc.add_paragraph()

    # ---- Metrics strip ----
    _table(doc, ["Clauses Extracted", "High-Risk Clauses", "Compliance Score", "Risk Level"], [[
        data.get("clauses_extracted", 0),
        data.get("high_risk_clause_count", 0),
        f"{data.get('compliance_score')}/100" if data.get("compliance_score") is not None else "—",
        data.get("risk_level", "Low"),
    ]])
    doc.add_paragraph()

    # ---- 2. Executive Summary ----
    _heading(doc, "2. Executive Summary")
    score = data.get("compliance_score")
    score_txt = f"{score}/100" if score is not None else "not computed"
    summary = (
        f"This report presents the results of a {data.get('task_type', 'legal').lower()} analysis. "
        f"The overall risk level is {data.get('risk_level', 'Low')}, based on "
        f"{data.get('high_risk_clause_count', 0)} high-severity finding(s) and a compliance score of {score_txt}."
    )
    doc.add_paragraph(summary)

    # ---- 3. Contract Analysis ----
    _heading(doc, "3. Contract Analysis")
    parties = data.get("parties") or []
    if parties:
        party_txt = ", ".join(
            (p.get("name", "") + (f" ({p['role']})" if p.get("role") else "")) if isinstance(p, dict) else str(p)
            for p in parties
        )
        doc.add_paragraph(f"Parties: {party_txt}")
    doc.add_paragraph(f"Effective date: {data.get('effective_date') or '—'}")
    doc.add_paragraph(f"Expiration date: {data.get('expiration_date') or '—'}")

    doc_intel = data.get("document_intelligence") or {}
    clause_rows = []
    for key, label in [
        ("payment_terms", "Payment Terms"), ("confidentiality_clause", "Confidentiality"),
        ("liability_clause", "Liability"), ("indemnity_clause", "Indemnity"),
        ("termination_clause", "Termination"),
    ]:
        detail = doc_intel.get(key) or {}
        status = "On file" if detail.get("present") else "Not found"
        clause_rows.append([label, status, detail.get("summary") or "—"])
    _table(doc, ["Clause", "Status", "Summary"], clause_rows)
    doc.add_paragraph()

    # ---- 4. Compliance Findings ----
    _heading(doc, "4. Compliance Findings")
    doc.add_paragraph(f"Compliance score: {data.get('compliance_score_value', 0)}/100")
    missing = data.get("missing_clauses") or []
    doc.add_paragraph(f"Missing clauses: {', '.join(missing) if missing else 'None'}")
    issues = data.get("compliance_issues") or []
    if issues:
        rows = [[i.get("issue", ""), i.get("severity", ""),
                  "; ".join(e.get("label", "") for e in (i.get("evidence") or [])) or "—"] for i in issues]
        _table(doc, ["Issue", "Severity", "Evidence"], rows)
    doc.add_paragraph()

    # ---- 5. Risk Assessment ----
    _heading(doc, "5. Risk Assessment")
    doc.add_paragraph(f"Overall risk level: {data.get('risk_level', 'Low')}")
    for r in data.get("risks") or []:
        doc.add_paragraph(f"[{r.get('risk_level','')}] {r.get('reason','')}", style="List Bullet")

    # ---- 6. High-Risk Clauses ----
    _heading(doc, "6. High-Risk Clauses")
    high_flags = data.get("high_risk_flags") or []
    if high_flags:
        rows = [[f.get("flag_type", ""), f.get("severity", ""), f.get("explanation", "")] for f in high_flags]
        _table(doc, ["Flag Type", "Severity", "Explanation"], rows)
    else:
        doc.add_paragraph("No high-severity red flags detected.")
    doc.add_paragraph()

    # ---- 7. Recommendations ----
    _heading(doc, "7. Recommendations")
    recs = data.get("recommendations") or []
    if recs:
        for r in recs:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("No specific review areas flagged.")

    # ---- 8. Evidence Sources ----
    _heading(doc, "8. Evidence Sources")
    evidence = data.get("evidence") or []
    if evidence:
        seen = set()
        rows = []
        for e in evidence:
            key = (e.get("label"), e.get("source"))
            if key in seen:
                continue
            seen.add(key)
            rows.append([e.get("label", ""), e.get("source", ""), e.get("quote") or "—"])
        _table(doc, ["Evidence", "Source", "Excerpt"], rows)
    else:
        doc.add_paragraph("No evidence references recorded.")
    sources = data.get("sources") or []
    if sources:
        doc.add_paragraph("Legal References Retrieved (RAG):")
        for s in sources[:10]:
            doc.add_paragraph(f"[{s.get('source_type','')}] {s.get('source_name','')}", style="List Bullet")

    # ---- 9. Agent Workflow Summary ----
    _heading(doc, "9. Agent Workflow Summary")
    doc.add_paragraph(f"Workflow version: {data.get('agent_workflow_version','')}")
    for i, stage in enumerate(PIPELINE_STAGES, start=1):
        doc.add_paragraph(f"\u2713 {i}. {stage} \u2014 complete")

    # ---- 10. Disclaimer + Timestamp ----
    doc.add_page_break()
    _heading(doc, "10. Disclaimer")
    p = doc.add_paragraph(DISCLAIMER_TEXT)
    for run in p.runs:
        run.font.color.rgb = MUTED
        run.font.size = Pt(9)
    doc.add_paragraph(f"Generated by: {data.get('generated_by')}")
    doc.add_paragraph(f"Generated on: {data.get('generated_on')}")
    doc.add_paragraph(f"Report ID: {data.get('report_id')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
